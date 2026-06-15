"""Document loaders for RAG pipeline."""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import structlog
from PyPDF2 import PdfReader

logger = structlog.get_logger()

EXTENSION_MAP: dict[str, str] = {
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
    ".text": "text",
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".csv": "csv",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".gif": "image",
    ".webp": "image",
}


@dataclass
class Document:
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)


def load_markdown(path: str) -> Document:
    text = Path(path).read_text(encoding="utf-8")
    logger.info("loaded_markdown", path=path, chars=len(text))
    return Document(content=text, metadata={"file_type": "markdown", "source": path})


def load_text(path: str) -> Document:
    text = Path(path).read_text(encoding="utf-8")
    logger.info("loaded_text", path=path, chars=len(text))
    return Document(content=text, metadata={"file_type": "text", "source": path})


def load_pdf(path: str) -> Document:
    reader = PdfReader(path)
    pages_text = []
    pages_to_ocr = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if len(text.strip()) < 50:
            pages_to_ocr.append(i)
        else:
            pages_text.append(text)

    if pages_to_ocr:
        logger.info("pdf_ocr_needed", path=path, pages=pages_to_ocr)
        ocr_text = _ocr_pdf_pages(path, pages_to_ocr)
        for i, text in zip(pages_to_ocr, ocr_text, strict=False):
            pages_text.insert(i, text)

    text = "\n".join(pages_text)
    is_scanned = len(pages_to_ocr) > len(reader.pages) / 2
    file_type = "pdf_scanned" if is_scanned else "pdf"

    logger.info(
        "loaded_pdf",
        path=path,
        pages=len(reader.pages),
        ocr_pages=len(pages_to_ocr),
        chars=len(text),
    )
    return Document(content=text, metadata={"file_type": file_type, "source": path})


def _ocr_pdf_pages(pdf_path: str, page_indices: list[int]) -> list[str]:
    """OCR specific pages from PDF using Tesseract"""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        images = convert_from_path(pdf_path, first_page=min(page_indices) + 1,
                                   last_page=max(page_indices) + 1)
        results = []
        for img in images:
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            results.append(text.strip())
        return results
    except Exception as e:
        logger.warning("ocr_failed", error=str(e))
        return [""] * len(page_indices)


def load_docx(path: str) -> Document:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    logger.info("loaded_docx", path=path, paragraphs=len(paragraphs), chars=len(text))
    return Document(content=text, metadata={"file_type": "docx", "source": path})


def load_xlsx(path: str) -> Document:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    rows: list[str] = []
    for sheet in wb.worksheets:
        rows.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip(" |"):
                rows.append(row_text)
    wb.close()
    text = "\n".join(rows)
    logger.info("loaded_xlsx", path=path, rows=len(rows), chars=len(text))
    return Document(content=text, metadata={"file_type": "xlsx", "source": path})


def load_csv(path: str) -> Document:
    import csv

    rows: list[str] = []
    with open(path, encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            row_text = " | ".join(row)
            if row_text.strip():
                rows.append(row_text)
    text = "\n".join(rows)
    logger.info("loaded_csv", path=path, rows=len(rows), chars=len(text))
    return Document(content=text, metadata={"file_type": "csv", "source": path})


def load_image(path: str) -> Document:
    logger.info("loaded_image", path=path)
    return Document(content="", metadata={"file_type": "image", "source": path})


_LOADERS = {
    "markdown": load_markdown,
    "text": load_text,
    "pdf": load_pdf,
    "docx": load_docx,
    "xlsx": load_xlsx,
    "csv": load_csv,
    "image": load_image,
}


def load_directory(path: str) -> list[Document]:
    docs: list[Document] = []
    for root, _dirs, files in os.walk(path):
        for fname in files:
            ext = Path(fname).suffix.lower()
            ftype = EXTENSION_MAP.get(ext)
            if ftype is None:
                logger.debug("skipping_unsupported", file=fname)
                continue
            fpath = os.path.join(root, fname)
            loader = _LOADERS[ftype]
            docs.append(loader(fpath))
    logger.info("loaded_directory", path=path, documents=len(docs))
    return docs
