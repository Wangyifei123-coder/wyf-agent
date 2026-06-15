"""Test document ingestion for all supported formats."""

import json
import os
import shutil
import sys
import tempfile
from pathlib import Path

import httpx

BASE_URL = 'http://localhost:8080'


def get_token():
    """Get authentication token."""
    r = httpx.post(f'{BASE_URL}/auth/login', json={'username': 'admin', 'password': 'admin123'})
    r.raise_for_status()
    data = r.json()
    if not data.get('token'):
        raise RuntimeError(f"Login failed: {data}")
    return data['token']


def create_test_docx(path: str):
    """Create a test DOCX file."""
    from docx import Document
    doc = Document()
    doc.add_heading('WYF Agent 技术架构', 0)
    doc.add_paragraph('WYF Agent 是一个企业级 AI 智能助手系统。')
    doc.add_paragraph('系统采用七层架构设计，包括：模型层、工具层、记忆层、推理层、编排层、安全层和可观测性层。')
    doc.add_paragraph('核心技术栈包括 FastAPI、ChromaDB、LangGraph 等。')
    doc.save(path)


def create_test_xlsx(path: str):
    """Create a test XLSX file."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = '技术栈'
    ws.append(['组件', '用途', '版本'])
    ws.append(['FastAPI', 'Web 框架', '0.115.0'])
    ws.append(['ChromaDB', '向量数据库', '0.6.0'])
    ws.append(['LangGraph', '推理引擎', '0.2.0'])
    ws.append(['SentenceTransformers', '文本嵌入', '3.4.0'])
    wb.save(path)


def create_test_csv(path: str):
    """Create a test CSV file."""
    with open(path, 'w', encoding='utf-8', newline='') as f:
        f.write('模块,功能,状态\n')
        f.write('Gateway,LLM 统一接口,已完成\n')
        f.write('Tools,工具注册与调用,已完成\n')
        f.write('Memory,短期/长期记忆管理,已完成\n')
        f.write('Reasoning,ReAct 推理循环,已完成\n')
        f.write('Orchestration,多 Agent 协作,已完成\n')
        f.write('Safety,输入过滤与输出审查,已完成\n')
        f.write('Observability,日志与追踪,已完成\n')


def create_test_pdf(path: str):
    """Create a test PDF file with actual text content."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(path, pagesize=letter)
        c.drawString(72, 700, "WYF Agent Technical Documentation")
        c.drawString(72, 680, "")
        c.drawString(72, 660, "WYF Agent is an enterprise-level AI assistant system.")
        c.drawString(72, 640, "It uses a seven-layer architecture: Gateway, Tools, Memory,")
        c.drawString(72, 620, "Reasoning, Orchestration, Safety, and Observability.")
        c.drawString(72, 600, "")
        c.drawString(72, 580, "Key technologies include FastAPI for the web framework,")
        c.drawString(72, 560, "ChromaDB for vector storage, and LangGraph for reasoning.")
        c.save()
    except ImportError:
        print("  [WARN] reportlab not installed, creating blank PDF")
        from PyPDF2 import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with open(path, 'wb') as f:
            writer.write(f)


def test_ingest_file(filepath: str, headers: dict):
    """Test ingesting a single file."""
    print(f"\n{'='*60}")
    print(f"Testing: {os.path.basename(filepath)}")
    print(f"{'='*60}")

    with tempfile.TemporaryDirectory() as tmpdir:
        dest = os.path.join(tmpdir, os.path.basename(filepath))
        shutil.copy2(filepath, dest)

        r = httpx.post(
            f'{BASE_URL}/knowledge/ingest',
            json={'path': tmpdir},
            headers=headers,
            timeout=30.0
        )
        r.raise_for_status()
        result = r.json()
        print(f"  Ingest result: {result}")
        return result


def test_query(question: str, headers: dict):
    """Test querying the knowledge base."""
    print(f"\n  Query: {question}")
    r = httpx.post(
        f'{BASE_URL}/chat',
        json={'message': question},
        headers=headers,
        timeout=180.0
    )
    if r.status_code != 200:
        print(f"  [WARN] Query returned status {r.status_code}: {r.text[:300]}")
        return None
    result = r.json()
    answer = result.get('answer', '')
    print(f"  Answer: {answer[:200]}...")
    return answer


def test_knowledge_sources(headers: dict):
    """Test getting knowledge sources."""
    r = httpx.get(f'{BASE_URL}/knowledge/sources', headers=headers)
    r.raise_for_status()
    return r.json()


def test_knowledge_stats(headers: dict):
    """Test getting knowledge stats."""
    r = httpx.get(f'{BASE_URL}/knowledge/stats', headers=headers)
    r.raise_for_status()
    return r.json()


def main():
    print("=" * 60)
    print("WYF Agent - Document Ingestion Test")
    print("=" * 60)

    try:
        r = httpx.get(f'{BASE_URL}/health', timeout=5.0)
        r.raise_for_status()
        print("[OK] Backend service is running")
    except Exception as e:
        print(f"[FAIL] Backend service is not running: {e}")
        print("Please start the backend service first: python -m src.api")
        sys.exit(1)

    token = get_token()
    headers = {'Authorization': f'Bearer {token}'}
    print("[OK] Authentication successful")

    test_docs_dir = os.path.join(os.path.dirname(__file__), 'data', 'test_docs')
    os.makedirs(test_docs_dir, exist_ok=True)

    results = {}

    md_files = list(Path(test_docs_dir).glob('*.md'))
    if md_files:
        for md_file in md_files:
            result = test_ingest_file(str(md_file), headers)
            results['markdown'] = result

    txt_files = list(Path(test_docs_dir).glob('*.txt'))
    if txt_files:
        for txt_file in txt_files:
            result = test_ingest_file(str(txt_file), headers)
            results['text'] = result

    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, 'test_document.pdf')
        create_test_pdf(pdf_path)
        result = test_ingest_file(pdf_path, headers)
        results['pdf'] = result

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'test_document.docx')
        create_test_docx(docx_path)
        result = test_ingest_file(docx_path, headers)
        results['docx'] = result

    with tempfile.TemporaryDirectory() as tmpdir:
        xlsx_path = os.path.join(tmpdir, 'test_document.xlsx')
        create_test_xlsx(xlsx_path)
        result = test_ingest_file(xlsx_path, headers)
        results['xlsx'] = result

    with tempfile.TemporaryDirectory() as tmpdir:
        csv_path = os.path.join(tmpdir, 'test_document.csv')
        create_test_csv(csv_path)
        result = test_ingest_file(csv_path, headers)
        results['csv'] = result

    print("\n" + "=" * 60)
    print("Ingestion Results Summary")
    print("=" * 60)
    for fmt, result in results.items():
        print(f"  {fmt}: {result}")

    print("\n" + "=" * 60)
    print("Verifying Ingestion via Knowledge Sources")
    print("=" * 60)
    sources_data = test_knowledge_sources(headers)
    source_list = sources_data.get('sources', [])
    print(f"  Total sources in knowledge base: {len(source_list)}")

    expected_formats = {
        'markdown': False,
        'text': False,
        'docx': False,
        'xlsx': False,
        'csv': False,
    }
    for src in source_list:
        src_lower = src.lower()
        if src_lower.endswith('.md'):
            expected_formats['markdown'] = True
        elif src_lower.endswith('.txt'):
            expected_formats['text'] = True
        elif src_lower.endswith('.docx'):
            expected_formats['docx'] = True
        elif src_lower.endswith('.xlsx'):
            expected_formats['xlsx'] = True
        elif src_lower.endswith('.csv'):
            expected_formats['csv'] = True

    print("\n  Format verification:")
    all_ok = True
    for fmt, found in expected_formats.items():
        status = "[OK]" if found else "[MISSING]"
        if not found:
            all_ok = False
        print(f"    {fmt}: {status}")

    if all_ok:
        print("\n  [OK] All formats successfully ingested and indexed!")
    else:
        print("\n  [WARN] Some formats are missing from the index")

    print("\n" + "=" * 60)
    print("Testing Query Functionality (requires LLM)")
    print("=" * 60)

    test_queries = [
        "What is WYF Agent?",
        "What are the components of WYF Agent's tech stack?",
        "What is the seven-layer architecture?",
        "What is FastAPI?",
        "What is ChromaDB used for?",
    ]

    query_results = []
    for query in test_queries:
        answer = test_query(query, headers)
        query_results.append((query, answer))

    print("\n" + "=" * 60)
    print("Knowledge Base Statistics")
    print("=" * 60)
    sources = test_knowledge_sources(headers)
    print(f"  Sources: {json.dumps(sources, ensure_ascii=False, indent=2)}")

    stats = test_knowledge_stats(headers)
    print(f"  Stats: {json.dumps(stats, ensure_ascii=False, indent=2)}")

    print("\n" + "=" * 60)
    print("Test Complete!")
    print("=" * 60)


if __name__ == '__main__':
    main()
